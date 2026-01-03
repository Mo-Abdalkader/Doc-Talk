import streamlit as st
from huggingface_hub import InferenceClient
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
import PyPDF2
import docx
import io
import tempfile

# Page configuration
st.set_page_config(
    page_title="Document Chat RAG System",
    page_icon="📚",
    layout="wide"
)

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_CHARS = 100000  # 100K characters
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'vectorstore' not in st.session_state:
    st.session_state.vectorstore = None
if 'document_processed' not in st.session_state:
    st.session_state.document_processed = False

def get_api_key():
    """Get API key from secrets or user input"""
    if 'HF_API_KEY' in st.secrets:
        return st.secrets['HF_API_KEY']
    return None

def configure_hf_client(api_key):
    """Configure HuggingFace client"""
    try:
        client = InferenceClient(token=api_key)
        return client
    except Exception as e:
        st.error(f"Error configuring HuggingFace API: {str(e)}")
        return None

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return None

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return None

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    try:
        return file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Error reading TXT: {str(e)}")
        return None

def process_document(file):
    """Process uploaded document and create vector store"""
    # Check file size
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    
    if file_size > MAX_FILE_SIZE:
        st.error(f"File size exceeds {MAX_FILE_SIZE / (1024*1024)}MB limit")
        return False
    
    # Extract text based on file type
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file)
    elif file_extension == 'docx':
        text = extract_text_from_docx(file)
    elif file_extension == 'txt':
        text = extract_text_from_txt(file)
    else:
        st.error("Unsupported file format. Please upload PDF, DOCX, or TXT files.")
        return False
    
    if not text:
        return False
    
    # Check text length
    if len(text) > MAX_CHARS:
        st.error(f"Document too large. Maximum {MAX_CHARS} characters allowed. Your document has {len(text)} characters.")
        return False
    
    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len
    )
    
    chunks = text_splitter.split_text(text)
    
    # Create documents
    documents = [Document(page_content=chunk) for chunk in chunks]
    
    # Create embeddings and vector store
    with st.spinner("Processing document... This may take a moment."):
        try:
            embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'}
            )
            
            vectorstore = FAISS.from_documents(documents, embeddings)
            st.session_state.vectorstore = vectorstore
            st.session_state.document_processed = True
            
            st.success(f"✅ Document processed! Created {len(chunks)} chunks for querying.")
            return True
            
        except Exception as e:
            st.error(f"Error creating vector store: {str(e)}")
            return False

def get_relevant_context(query, k=3):
    """Retrieve relevant context from vector store"""
    if st.session_state.vectorstore is None:
        return None
    
    try:
        docs = st.session_state.vectorstore.similarity_search(query, k=k)
        context = "\n\n".join([doc.page_content for doc in docs])
        return context
    except Exception as e:
        st.error(f"Error retrieving context: {str(e)}")
        return None

def generate_response(query, context, api_key):
    """Generate response using HuggingFace"""
    try:
        client = InferenceClient(token=api_key)
        
        messages = [
            {
                "role": "user",
                "content": f"""Based on the following context from the document, answer the user's question. 
If the answer cannot be found in the context, say so clearly.

Context:
{context}

Question: {query}

Answer:"""
            }
        ]
        
        # Using Llama 3.2 3B Instruct with chat completion
        response = client.chat_completion(
            messages=messages,
            model="meta-llama/Llama-3.2-3B-Instruct",
            max_tokens=500,
            temperature=0.7,
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"Error generating response: {str(e)}")
        return None

# Main UI
st.title("📚 Document Chat RAG System")
st.markdown("Upload a document and ask questions about its content!")

# Sidebar for API key and document upload
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # API Key input
    api_key = get_api_key()
    if not api_key:
        api_key = st.text_input(
            "Enter HuggingFace API Key",
            type="password",
            help="Get your free API key from https://huggingface.co/settings/tokens"
        )
    else:
        st.success("✅ API Key loaded from secrets")
    
    if api_key:
        client = configure_hf_client(api_key)
        if client:
            st.success("✅ HuggingFace API configured")
    
    st.divider()
    
    # Document upload
    st.header("📄 Upload Document")
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['pdf', 'docx', 'txt'],
        help=f"Max size: {MAX_FILE_SIZE/(1024*1024)}MB, Max length: {MAX_CHARS} characters"
    )
    
    if uploaded_file and not st.session_state.document_processed:
        if st.button("Process Document", type="primary"):
            if not api_key:
                st.error("Please enter your HuggingFace API key first!")
            else:
                process_document(uploaded_file)
    
    if st.session_state.document_processed:
        st.success("✅ Document ready for questions!")
        if st.button("Upload New Document"):
            st.session_state.messages = []
            st.session_state.vectorstore = None
            st.session_state.document_processed = False
            st.rerun()
    
    st.divider()
    
    # Information
    st.header("ℹ️ About")
    st.markdown("""
    This RAG system allows you to:
    - Upload PDF, DOCX, or TXT files
    - Ask questions about the content
    - Get AI-powered answers
    
    **Constraints:**
    - Max file size: 10MB
    - Max length: 100K chars
    - Supported: PDF, DOCX, TXT
    """)

# Main chat interface
if not api_key:
    st.warning("👈 Please enter your HuggingFace API key in the sidebar to get started.")
elif not st.session_state.document_processed:
    st.info("👈 Please upload a document in the sidebar to begin chatting.")
else:
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask a question about your document..."):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                context = get_relevant_context(prompt)
                
                if context:
                    response = generate_response(prompt, context, api_key)
                    
                    if response:
                        st.markdown(response)
                        st.session_state.messages.append({"role": "assistant", "content": response})
                    else:
                        st.error("Failed to generate response.")
                else:
                    st.error("Failed to retrieve relevant context.")

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px;'>
    <small>Built with Streamlit, LangChain, FAISS, and HuggingFace Llama 3.2 | 100% Free to Use</small>
</div>
""", unsafe_allow_html=True)
